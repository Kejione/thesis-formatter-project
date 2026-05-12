"""
Format check and fix Celery tasks.

Implements the full processing pipeline:
1. process_format_check  - Download thesis, resolve rules, run checker, save issues.
2. process_format_fix    - Download thesis, apply fixes, upload fixed doc, save changes.
3. parse_spec_file       - Download spec, extract text, AI-parse rules, validate & persist.
"""

import asyncio
import os
import tempfile
import uuid
from typing import Optional

from loguru import logger

from app.core.database import get_db_context
from app.models import Change, Issue, Rule, Task
from app.services.ai import SpecParser
from app.services.ai.provider import get_model_manager
from app.services.docx import DocumentProcessor
from app.services.docx.generator import DocxGenerator
from app.services.rule import get_rule_engine
from app.services.storage import get_storage_service
from app.tasks.celery_app import celery_app


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _download_to_temp(object_key: str, suffix: str = ".docx") -> str:
    """Download a file from MinIO and write it to a temporary path.

    Returns the temporary file path.  The caller is responsible for
    cleaning up the file after use.
    """
    storage = get_storage_service()
    file_data = storage.download_file(object_key)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(file_data)
    tmp.close()
    logger.debug("文件已下载到临时路径: {} ({} bytes)", tmp.name, len(file_data))
    return tmp.name


def _cleanup_temp(path: str) -> None:
    """Safely remove a temporary file."""
    try:
        if path and os.path.exists(path):
            os.unlink(path)
            logger.debug("临时文件已清理: {}", path)
    except OSError as exc:
        logger.warning("清理临时文件失败 {}: {}", path, exc)


async def _get_task(db, task_id: str) -> Task:
    """Fetch a Task by UUID string, raising if not found."""
    from sqlalchemy import select

    result = await db.execute(select(Task).where(Task.id == uuid.UUID(task_id)))
    task = result.scalar_one_or_none()
    if task is None:
        raise ValueError(f"Task {task_id} 不存在")
    return task


async def _update_task_status(
    task_id: str,
    status: str,
    *,
    result_summary: Optional[dict] = None,
    fixed_file_key: Optional[str] = None,
    error_message: Optional[str] = None,
    rule_snapshot: Optional[dict] = None,
) -> None:
    """Update task fields and commit inside an async DB context."""
    async with get_db_context() as db:
        task = await _get_task(db, task_id)
        task.status = status
        if result_summary is not None:
            task.result_summary = result_summary
        if fixed_file_key is not None:
            task.fixed_file_key = fixed_file_key
        if error_message is not None:
            task.error_message = error_message
        if rule_snapshot is not None:
            task.rule_snapshot = rule_snapshot
        await db.commit()
        logger.info("Task {} 状态更新为 {}", task_id, status)


def _extract_spec_text(file_path: str) -> str:
    """Extract text from a specification file (.pdf, .docx, .txt).

    Returns the extracted plain-text string.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    if ext == ".pdf":
        import pdfplumber

        pages_text: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
        return "\n\n".join(pages_text)

    if ext == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    raise ValueError(f"不支持的规范文件格式: {ext}")


# ---------------------------------------------------------------------------
# Task 1: process_format_check
# ---------------------------------------------------------------------------

@celery_app.task(bind=True, name="process_format_check")
def process_format_check(self, task_id: str):
    """
    执行论文格式检查。

    流程:
    1. 更新任务状态为 processing
    2. 从 MinIO 下载论文文件
    3. 解析规则来源（spec_file / template / 默认）
    4. 保存 rule_snapshot
    5. 调用 DocumentProcessor.process(fix=False) 检查格式
    6. 将问题写入数据库
    7. 更新任务状态为 completed
    """
    logger.info("开始格式检查任务: {}", task_id)
    temp_path: Optional[str] = None

    try:
        # 1. 更新状态
        asyncio.run(_update_task_status(task_id, "processing"))

        # 2. 获取任务信息 & 下载论文
        async def _prepare():
            async with get_db_context() as db:
                task = await _get_task(db, task_id)
                thesis_key = task.thesis_file_key
                spec_key = task.spec_file_key
                template_id = task.template_id
                model_id = task.model_id
                return thesis_key, spec_key, template_id, model_id

        thesis_key, spec_key, template_id, model_id = asyncio.run(_prepare())
        temp_path = _download_to_temp(thesis_key, suffix=".docx")

        # 3. 解析规则
        rules = _resolve_rules(task_id, spec_key, template_id, model_id)

        # 4. 保存 rule_snapshot
        asyncio.run(_update_task_status(task_id, "processing", rule_snapshot=rules))

        # 5. 执行格式检查
        processor = DocumentProcessor(rules)
        result = processor.process(temp_path, fix=False, generate_reports=False)

        # 6. 将 issues 写入数据库
        async def _save_issues():
            async with get_db_context() as db:
                for issue in result.issues:
                    db_issue = Issue(
                        task_id=uuid.UUID(task_id),
                        severity=issue.severity.value,
                        category=issue.category.value,
                        location=issue.location,
                        rule_id=issue.rule_id,
                        current_value=issue.current_value,
                        expected_value=issue.expected_value,
                        suggestion=issue.suggestion,
                        is_fixed=False,
                    )
                    db.add(db_issue)
                await db.commit()
                logger.info("Task {} 共保存 {} 条问题记录", task_id, len(result.issues))

        asyncio.run(_save_issues())

        # 7. 构建结果摘要 & 更新状态
        error_count = sum(1 for i in result.issues if i.severity.value == "error")
        warning_count = sum(1 for i in result.issues if i.severity.value == "warning")
        info_count = sum(1 for i in result.issues if i.severity.value == "info")

        summary = {
            "total_issues": len(result.issues),
            "error_count": error_count,
            "warning_count": warning_count,
            "info_count": info_count,
        }

        asyncio.run(_update_task_status(task_id, "completed", result_summary=summary))

        logger.info(
            "格式检查完成: task={} errors={} warnings={} info={}",
            task_id, error_count, warning_count, info_count,
        )
        return {"status": "completed", "task_id": task_id, "summary": summary}

    except Exception as exc:
        logger.error("格式检查失败 task={}: {}", task_id, exc)
        try:
            asyncio.run(_update_task_status(task_id, "failed", error_message=str(exc)))
        except Exception as inner_exc:
            logger.error("更新失败状态时出错: {}", inner_exc)
        raise

    finally:
        _cleanup_temp(temp_path)


def _resolve_rules(
    task_id: str,
    spec_key: Optional[str],
    template_id: Optional[uuid.UUID],
    model_id: Optional[str],
) -> dict:
    """根据优先级解析格式规则: spec_file > template > 默认。

    当 spec_file_key 存在时，先同步调用 parse_spec_file 的核心逻辑
    来获取 AI 解析的规则。若解析失败则回退到默认规则。
    """
    rule_engine = get_rule_engine()

    # 优先级 1: 规范文件 (AI 解析)
    if spec_key:
        logger.info("检测到规范文件，尝试 AI 解析规则: {}", spec_key)
        try:
            rules = _parse_spec_rules_inline(spec_key, model_id)
            if rules:
                # 用默认规则做底，深度合并 AI 解析结果
                default_rules = rule_engine.get_default_rules()
                rules = rule_engine.merge_rules(default_rules, rules)
                logger.info("AI 规范解析成功，共获取规则字段: {}", len(rules))
                return rules
        except Exception as exc:
            logger.warning("AI 规范解析失败，将回退到其他规则来源: {}", exc)

    # 优先级 2: 模板规则
    if template_id:
        logger.info("加载模板规则: template_id={}", template_id)
        try:
            rules = asyncio.run(_load_template_rules(task_id, template_id))
            if rules:
                return rules
        except Exception as exc:
            logger.warning("加载模板规则失败，将使用默认规则: {}", exc)

    # 优先级 3: 默认规则
    logger.info("使用默认格式规则")
    return rule_engine.get_default_rules()


async def _load_template_rules(task_id: str, template_id: uuid.UUID) -> Optional[dict]:
    """从数据库加载模板关联的规则。"""
    from sqlalchemy import select
    from app.models import Template

    async with get_db_context() as db:
        result = await db.execute(
            select(Template).where(Template.id == template_id)
        )
        template = result.scalar_one_or_none()
        if template and template.rule:
            rule_engine = get_rule_engine()
            default_rules = rule_engine.get_default_rules()
            return rule_engine.merge_rules(default_rules, template.rule.rule_data)
    return None


def _parse_spec_rules_inline(spec_file_key: str, model_id: Optional[str]) -> Optional[dict]:
    """同步下载规范文件并调用 AI 解析，返回规则字典。"""
    temp_path: Optional[str] = None
    try:
        temp_path = _download_to_temp(spec_file_key, suffix=_guess_suffix(spec_file_key))
        text = _extract_spec_text(temp_path)

        if not text.strip():
            logger.warning("规范文件内容为空: {}", spec_file_key)
            return None

        model_manager = get_model_manager()
        parser = SpecParser(model_manager)
        rules = asyncio.run(parser.parse(text=text, model_id=model_id))
        return rules
    finally:
        _cleanup_temp(temp_path)


def _guess_suffix(key: str) -> str:
    """根据 object_key 猜测文件后缀。"""
    lower = key.lower()
    for ext in (".pdf", ".docx", ".txt"):
        if lower.endswith(ext):
            return ext
    return ".bin"


# ---------------------------------------------------------------------------
# Task 2: process_format_fix
# ---------------------------------------------------------------------------

@celery_app.task(bind=True, name="process_format_fix")
def process_format_fix(self, task_id: str, issue_ids: Optional[list[str]] = None):
    """
    执行论文格式修复。

    流程:
    1. 更新任务状态为 fixing
    2. 从 MinIO 下载论文
    3. 从 task.rule_snapshot 获取规则
    4. 调用 DocumentProcessor.fix_only() 修复
    5. 上传修复后的文档到 MinIO
    6. 保存修改记录到数据库
    7. 生成并保存 change_log / report 到 MinIO
    8. 更新任务状态为 fixed
    """
    logger.info("开始格式修复任务: task={} issue_ids={}", task_id, issue_ids)
    temp_path: Optional[str] = None

    try:
        # 1. 更新状态
        asyncio.run(_update_task_status(task_id, "fixing"))

        # 2. 获取任务信息 & 下载论文
        async def _prepare():
            async with get_db_context() as db:
                task = await _get_task(db, task_id)
                return task.thesis_file_key, task.rule_snapshot

        thesis_key, rule_snapshot = asyncio.run(_prepare())

        if not rule_snapshot:
            raise ValueError("任务缺少 rule_snapshot，无法执行修复")

        temp_path = _download_to_temp(thesis_key, suffix=".docx")

        # 3. 执行修复
        processor = DocumentProcessor(rule_snapshot)
        changes, fixed_bytes = processor.fix_only(temp_path, issue_ids=issue_ids)

        # 4. 上传修复后的文档
        storage = get_storage_service()
        fixed_key = f"fixed/{task_id}.docx"
        storage.upload_file(fixed_bytes, fixed_key, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        logger.info("修复后文档已上传: {}", fixed_key)

        # 5. 保存修改记录到数据库
        async def _save_changes():
            async with get_db_context() as db:
                for change in changes:
                    db_change = Change(
                        task_id=uuid.UUID(task_id),
                        issue_id=None,  # change 记录通过 rule_id 关联
                        category=change.category,
                        location=change.location,
                        before_value=change.before_value,
                        after_value=change.after_value,
                        risk_level=change.risk_level,
                    )
                    db.add(db_change)
                await db.commit()
                logger.info("Task {} 共保存 {} 条修改记录", task_id, len(changes))

        asyncio.run(_save_changes())

        # 6. 生成报告并上传到 MinIO
        _generate_and_upload_reports(task_id, temp_path, rule_snapshot, changes)

        # 7. 更新任务状态
        asyncio.run(
            _update_task_status(
                task_id,
                "fixed",
                fixed_file_key=fixed_key,
                result_summary={
                    "total_changes": len(changes),
                    "low_risk": sum(1 for c in changes if c.risk_level == "low"),
                    "medium_risk": sum(1 for c in changes if c.risk_level == "medium"),
                    "high_risk": sum(1 for c in changes if c.risk_level == "high"),
                },
            )
        )

        logger.info("格式修复完成: task={} changes={}", task_id, len(changes))
        return {"status": "fixed", "task_id": task_id, "total_changes": len(changes)}

    except Exception as exc:
        logger.error("格式修复失败 task={}: {}", task_id, exc)
        try:
            asyncio.run(_update_task_status(task_id, "failed", error_message=str(exc)))
        except Exception as inner_exc:
            logger.error("更新失败状态时出错: {}", inner_exc)
        raise

    finally:
        _cleanup_temp(temp_path)


def _generate_and_upload_reports(
    task_id: str,
    original_path: str,
    rules: dict,
    changes: list,
) -> None:
    """生成 change_log_markdown 和 report_markdown 并上传到 MinIO。"""
    if not changes:
        logger.info("无修改记录，跳过报告生成: task={}", task_id)
        return

    from docx import Document as DocxDocument

    document = DocxDocument(original_path)
    generator = DocxGenerator(document, os.path.basename(original_path))

    doc_info_dict = {"文件名": os.path.basename(original_path)}

    change_log_md = generator.generate_change_log_markdown(changes, doc_info_dict)
    report_md = generator.generate_report_markdown([], changes, doc_info_dict, rules)

    storage = get_storage_service()

    change_log_key = f"reports/{task_id}/change_log.md"
    report_key = f"reports/{task_id}/report.md"

    storage.upload_file(
        change_log_md.encode("utf-8"),
        change_log_key,
        content_type="text/markdown",
    )
    storage.upload_file(
        report_md.encode("utf-8"),
        report_key,
        content_type="text/markdown",
    )

    logger.info("报告已上传: {} / {}", change_log_key, report_key)


# ---------------------------------------------------------------------------
# Task 3: parse_spec_file
# ---------------------------------------------------------------------------

@celery_app.task(bind=True, name="parse_spec_file")
def parse_spec_file(
    self,
    task_id: str,
    spec_file_key: str,
    model_id: Optional[str] = None,
):
    """
    使用 AI 解析格式规范文件。

    流程:
    1. 从 MinIO 下载规范文件
    2. 提取文本内容 (PDF / DOCX / TXT)
    3. 调用 SpecParser.parse() 解析规则
    4. 使用 RuleEngine.validate_rules() 校验
    5. 将规则保存为 Rule 记录 (source="ai_parsed")
    6. 更新 task.rule_snapshot
    """
    logger.info("开始解析规范文件: task={} spec={}", task_id, spec_file_key)
    temp_path: Optional[str] = None

    try:
        # 1. 下载规范文件
        temp_path = _download_to_temp(spec_file_key, suffix=_guess_suffix(spec_file_key))

        # 2. 提取文本
        text = _extract_spec_text(temp_path)
        if not text.strip():
            raise ValueError("规范文件内容为空，无法解析")

        logger.info("规范文件文本提取完成，共 {} 字符", len(text))

        # 3. 调用 AI 解析
        model_manager = get_model_manager()
        parser = SpecParser(model_manager)
        rules = asyncio.run(parser.parse(text=text, model_id=model_id))

        logger.info("AI 解析完成，获取规则字段: {}", list(rules.keys()))

        # 4. 校验规则
        rule_engine = get_rule_engine()
        validation_errors = rule_engine.validate_rules(rules)

        if validation_errors:
            logger.warning("规则校验发现 {} 个问题: {}", len(validation_errors), validation_errors)
            # 不中断流程，仍然保存规则，但记录警告

        # 5. 保存规则到数据库
        async def _save_rule():
            async with get_db_context() as db:
                school_name = rules.get("school_name", "未知学校")
                rule_record = Rule(
                    name=f"AI 解析规则 - {school_name}",
                    source="ai_parsed",
                    rule_data=rules,
                    school_name=school_name,
                    is_active=True,
                )
                db.add(rule_record)
                await db.flush()  # 获取 rule_record.id

                # 6. 更新 task.rule_snapshot
                task = await _get_task(db, task_id)
                task.rule_snapshot = rules
                await db.commit()

                logger.info(
                    "规则已保存: rule_id={} task_id={}",
                    rule_record.id,
                    task_id,
                )
                return rule_record.id

        rule_id = asyncio.run(_save_rule())

        logger.info("规范文件解析完成: task={} rule_id={}", task_id, rule_id)
        return {"status": "parsed", "task_id": task_id, "rule_id": str(rule_id)}

    except Exception as exc:
        logger.error("规范文件解析失败 task={}: {}", task_id, exc)
        try:
            asyncio.run(
                _update_task_status(task_id, "failed", error_message=str(exc))
            )
        except Exception as inner_exc:
            logger.error("更新失败状态时出错: {}", inner_exc)
        raise

    finally:
        _cleanup_temp(temp_path)
