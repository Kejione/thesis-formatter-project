"""
Pytest configuration and fixtures for document processing tests.
"""

import pytest
import tempfile
import os
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Pytest configuration
pytest_plugins = ["pytest_asyncio"]
pytest_asyncio_mode = "auto"


@pytest.fixture(scope="session")
def event_loop():
    """Create an instance of the default event loop for async tests."""
    import asyncio
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest.fixture
def temp_docx_file():
    """
    Create a temporary DOCX file with sample content for testing.
    
    Returns a path to a temporary DOCX file that is automatically cleaned up.
    """
    doc = Document()
    
    # Set document margins (non-standard for testing)
    section = doc.sections[0]
    section.top_margin = Cm(2.0)  # Should be 2.5cm
    section.bottom_margin = Cm(2.0)  # Should be 2.5cm
    section.left_margin = Cm(2.5)  # Should be 3.0cm
    section.right_margin = Cm(2.5)  # Should be 2.5cm
    
    # Add title heading
    title = doc.add_heading("毕业论文测试文档", level=1)
    
    # Add abstract section
    abstract_heading = doc.add_heading("摘要", level=2)
    abstract_para = doc.add_paragraph(
        "这是一篇测试论文的摘要部分。本研究旨在测试文档处理模块的功能。"
        "通过构建测试用例，验证解析器、检查器和修复器的正确性。"
    )
    # Set font properties for testing
    for run in abstract_para.runs:
        run.font.size = Pt(10.5)  # Wrong size (should be 12pt)
        run.font.name = "Arial"  # Wrong font (should be Times New Roman)
        # Set East Asian font
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), "微软雅黑")  # Wrong Chinese font
    
    # Set paragraph format
    abstract_para.paragraph_format.line_spacing = 1.0  # Wrong spacing (should be 1.5)
    abstract_para.paragraph_format.first_line_indent = Cm(0.5)  # Wrong indent
    
    # Add introduction section
    intro_heading = doc.add_heading("第一章 绪论", level=1)
    intro_para1 = doc.add_paragraph(
        "这是第一章的内容。文档处理是毕业论文排版的重要环节。"
        "正确的格式设置能够提升论文的专业性和可读性。"
    )
    for run in intro_para1.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), "宋体")
    intro_para1.paragraph_format.line_spacing = 1.5
    
    # Add another paragraph with wrong spacing
    intro_para2 = doc.add_paragraph(
        "这是第二个段落，用于测试段落间距的设置。"
    )
    for run in intro_para2.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), "宋体")
    intro_para2.paragraph_format.line_spacing = 2.0  # Wrong: should be 1.5
    intro_para2.paragraph_format.space_before = Pt(12)
    intro_para2.paragraph_format.space_after = Pt(12)
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"单元格 {i+1},{j+1}"
    
    # Add caption before table
    table_caption = doc.add_paragraph("表1-1 测试数据表")
    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add references section
    ref_heading = doc.add_heading("参考文献", level=1)
    ref1 = doc.add_paragraph(
        "[1] 张三, 李四. 文档处理技术研究[J]. 计算机学报, 2023, 45(3): 100-120."
    )
    ref1.paragraph_format.first_line_indent = Cm(0)  # Wrong: should have hanging indent
    
    ref2 = doc.add_paragraph(
        "[2] Wang W, Zhang L. Advanced Document Processing[M]. Springer, 2022."
    )
    ref2.paragraph_format.first_line_indent = Cm(0)
    
    # Save to temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name
    
    yield temp_path
    
    # Cleanup
    if os.path.exists(temp_path):
        os.unlink(temp_path)


@pytest.fixture
def empty_docx_file():
    """Create an empty DOCX file for edge case testing."""
    doc = Document()
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name
    
    yield temp_path
    
    if os.path.exists(temp_path):
        os.unlink(temp_path)


@pytest.fixture
def sample_rules():
    """
    Return a dictionary with format rules for testing.
    
    These rules represent typical Chinese thesis formatting requirements.
    """
    return {
        "page_margin": {
            "top": "2.5cm",
            "bottom": "2.5cm",
            "left": "3.0cm",
            "right": "2.5cm",
        },
        "font": {
            "cn_body": "宋体",
            "en_body": "Times New Roman",
            "cn_heading": "黑体",
            "en_heading": "Arial",
        },
        "font_size": {
            "body": "12pt",
            "heading1": "22pt",
            "heading2": "16pt",
            "heading3": "14pt",
        },
        "line_spacing": {
            "body": "1.5倍",
            "heading": "单倍",
        },
        "paragraph_spacing": {
            "body": {
                "before": "0pt",
                "after": "0pt",
            },
        },
        "indent": {
            "first_line": "2字符",
        },
        "heading_style": {
            "heading1": {
                "font": "黑体",
                "bold": True,
            },
            "heading2": {
                "font": "黑体",
                "bold": True,
            },
        },
        "page_number": {
            "position": "bottom_center",
            "format": "arabic",
        },
        "toc": {
            "required": True,
        },
        "references": {
            "indent": "hanging",
        },
    }


@pytest.fixture
def minimal_rules():
    """Return minimal rules for basic testing."""
    return {
        "page_margin": {
            "top": "2.54cm",
            "bottom": "2.54cm",
            "left": "3.17cm",
            "right": "3.17cm",
        },
    }


@pytest.fixture
def complex_docx_file():
    """
    Create a complex DOCX file with various elements for comprehensive testing.
    """
    doc = Document()
    
    # Set section properties
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    
    # Add cover page content
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("测试论文标题")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), "黑体")
    
    # Add page break
    doc.add_page_break()
    
    # Add TOC placeholder
    toc_heading = doc.add_heading("目录", level=1)
    toc_para = doc.add_paragraph("目录内容将在生成时插入...")
    
    # Add main content
    for i in range(1, 4):
        heading = doc.add_heading(f"第{i}章 测试章节", level=1)
        
        for j in range(1, 3):
            subheading = doc.add_heading(f"{i}.{j} 测试小节", level=2)
            
            para = doc.add_paragraph(
                f"这是第{i}章第{j}节的内容。用于测试文档解析和格式检查功能。"
                "包含多个段落以测试分页和页码计算。"
            )
            for run in para.runs:
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), "宋体")
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = Cm(0.74)  # ~2 chars
    
    # Add images section
    img_heading = doc.add_heading("图表清单", level=1)
    
    # Add references
    ref_heading = doc.add_heading("参考文献", level=1)
    for i in range(1, 4):
        ref = doc.add_paragraph(
            f"[{i}] 作者{i}. 测试文献标题{i}[J]. 测试期刊, 2023, {i}(1): {i}0-{i}9."
        )
        # Set hanging indent
        ref.paragraph_format.first_line_indent = Pt(-24)
        ref.paragraph_format.left_indent = Pt(24)
    
    # Save to temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name
    
    yield temp_path
    
    if os.path.exists(temp_path):
        os.unlink(temp_path)
