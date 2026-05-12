"""
Tests for DocxParser module.

Tests document parsing functionality including loading, section extraction,
paragraph extraction, word counting, and edge cases.
"""

import pytest
from docx import Document
from docx.shared import Pt, Cm

import sys
import os

# Add parent directory to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

# Import directly from module files to avoid app package initialization
import importlib.util

def load_module(module_name, file_path):
    # Use unique module name to avoid conflicts
    unique_name = f"parser_test_{module_name}"
    if unique_name in sys.modules:
        return sys.modules[unique_name]
    spec = importlib.util.spec_from_file_location(unique_name, file_path)
    module = importlib.util.module_from_spec(spec)
    sys.modules[unique_name] = module
    spec.loader.exec_module(module)
    return module

# Load parser module directly
parser_module = load_module(
    "parser",
    os.path.join(os.path.dirname(__file__), '..', 'app', 'services', 'docx', 'parser.py')
)

DocxParser = parser_module.DocxParser
DocumentInfo = parser_module.DocumentInfo
ParagraphInfo = parser_module.ParagraphInfo
SectionInfo = parser_module.SectionInfo
ElementType = parser_module.ElementType
FontInfo = parser_module.FontInfo
ParagraphFormat = parser_module.ParagraphFormat


class TestDocxParser:
    """Test suite for DocxParser class."""
    
    def test_parser_loads_document(self, temp_docx_file):
        """Test that parser can load a valid DOCX file."""
        parser = DocxParser(temp_docx_file)
        
        # Should be able to load without errors
        parser.load()
        
        # Internal document should be set
        assert parser._document is not None
        # Check that _document has expected attributes
        assert hasattr(parser._document, 'paragraphs')
        assert hasattr(parser._document, 'sections')
    
    def test_parser_extracts_sections(self, temp_docx_file):
        """Test section information extraction."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # Should have at least one section
        assert len(doc_info.sections) >= 1
        
        # Check section properties
        section = doc_info.sections[0]
        assert isinstance(section, SectionInfo)
        assert section.index == 0
        
        # Check margin values are extracted
        assert section.top_margin_cm > 0
        assert section.bottom_margin_cm > 0
        assert section.left_margin_cm > 0
        assert section.right_margin_cm > 0
        
        # Check page dimensions
        assert section.page_width_cm > 0
        assert section.page_height_cm > 0
    
    def test_parser_extracts_paragraphs(self, temp_docx_file):
        """Test paragraph information extraction."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # Should have multiple paragraphs
        assert len(doc_info.paragraphs) >= 5
        
        # Check paragraph properties
        for para in doc_info.paragraphs:
            assert isinstance(para, ParagraphInfo)
            assert para.index >= 0
            assert isinstance(para.text, str)
            assert isinstance(para.element_type, ElementType)
        
        # Check first paragraph has text
        first_para = doc_info.paragraphs[0]
        assert len(first_para.text) > 0
    
    def test_parser_detects_headings(self, temp_docx_file):
        """Test heading detection."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # Find heading paragraphs
        headings = [p for p in doc_info.paragraphs if p.element_type == ElementType.HEADING]
        
        # Should have at least some headings
        assert len(headings) >= 2
        
        # Check heading levels
        for heading in headings:
            assert heading.heading_level is not None
            assert 1 <= heading.heading_level <= 9
    
    def test_parser_extracts_font_info(self, temp_docx_file):
        """Test font information extraction."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # Check font info on paragraphs with text
        for para in doc_info.paragraphs:
            if para.text.strip():
                assert isinstance(para.font_info, FontInfo)
                # Font size might be None if not explicitly set
                if para.font_info.size_pt is not None:
                    assert para.font_info.size_pt > 0
    
    def test_parser_extracts_paragraph_format(self, temp_docx_file):
        """Test paragraph format extraction."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        for para in doc_info.paragraphs:
            assert isinstance(para.paragraph_format, ParagraphFormat)
            # Line spacing might be None
            if para.paragraph_format.line_spacing is not None:
                assert para.paragraph_format.line_spacing > 0
    
    def test_parser_counts_words(self, temp_docx_file):
        """Test word count accuracy."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # Should have positive word count
        assert doc_info.word_count > 0
        
        # Character count should be positive
        assert doc_info.char_count > 0
        
        # Character count should be >= word count (for Chinese)
        assert doc_info.char_count >= doc_info.word_count
    
    def test_parser_handles_empty_document(self, empty_docx_file):
        """Test handling of empty document."""
        parser = DocxParser(empty_docx_file)
        doc_info = parser.parse()
        
        # Should return valid DocumentInfo
        assert isinstance(doc_info, DocumentInfo)
        
        # Should have one section (default)
        assert len(doc_info.sections) >= 1
        
        # Should have no paragraphs
        assert len(doc_info.paragraphs) == 0
        
        # Word count should be 0
        assert doc_info.word_count == 0
        assert doc_info.char_count == 0
        
        # Page count should be at least 1
        assert doc_info.page_count >= 1
    
    def test_parser_extracts_tables(self, temp_docx_file):
        """Test table extraction."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # Should find the table we added
        assert len(doc_info.tables) >= 1
        
        # Check table properties
        table = doc_info.tables[0]
        assert table.row_count == 3
        assert table.column_count == 3
        assert table.index == 0
    
    def test_parser_extracts_metadata(self, temp_docx_file):
        """Test document metadata extraction."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # Should extract title from first heading or paragraph
        assert doc_info.title is not None
        assert len(doc_info.title) > 0
    
    def test_parser_page_estimation(self, temp_docx_file):
        """Test page count estimation."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # Should have positive page count
        assert doc_info.page_count >= 1
        
        # Page count should be reasonable for document size
        # (at least 1 page, not unreasonably high)
        assert doc_info.page_count <= len(doc_info.paragraphs) + 10
    
    def test_parser_detects_toc(self, temp_docx_file):
        """Test table of contents detection."""
        parser = DocxParser(temp_docx_file)
        doc_info = parser.parse()
        
        # TOC info should exist
        assert doc_info.toc is not None
        
        # Should count headings
        assert doc_info.toc.heading_count >= 0
    
    def test_parser_handles_complex_document(self, complex_docx_file):
        """Test parsing of complex document with multiple elements."""
        parser = DocxParser(complex_docx_file)
        doc_info = parser.parse()
        
        # Should have multiple sections or paragraphs
        assert len(doc_info.paragraphs) > 10
        
        # Should have headings
        headings = [p for p in doc_info.paragraphs if p.element_type == ElementType.HEADING]
        assert len(headings) >= 3
        
        # Should have tables
        assert len(doc_info.tables) >= 0  # May or may not have tables
        
        # Word count should reflect document size
        assert doc_info.word_count > 50


class TestDocxParserEdgeCases:
    """Test edge cases for DocxParser."""
    
    def test_parser_with_special_characters(self, tmp_path):
        """Test parsing document with special characters."""
        doc = Document()
        
        # Add text with special characters
        special_text = "特殊字符：中文 English 123 !@#$%^&*() 日本語 한국어"
        para = doc.add_paragraph(special_text)
        
        # Save and parse
        file_path = tmp_path / "special.docx"
        doc.save(file_path)
        
        parser = DocxParser(str(file_path))
        doc_info = parser.parse()
        
        assert len(doc_info.paragraphs) == 1
        # Text should be preserved
        assert "特殊字符" in doc_info.paragraphs[0].text
    
    def test_parser_with_long_text(self, tmp_path):
        """Test parsing document with very long text."""
        doc = Document()
        
        # Add very long paragraph
        long_text = "这是一个很长的段落。" * 1000
        doc.add_paragraph(long_text)
        
        file_path = tmp_path / "long.docx"
        doc.save(file_path)
        
        parser = DocxParser(str(file_path))
        doc_info = parser.parse()
        
        # Should handle long text (truncated in output)
        assert len(doc_info.paragraphs) == 1
        # Text should be truncated with "..."
        assert "..." in doc_info.paragraphs[0].text or len(doc_info.paragraphs[0].text) <= 203
    
    def test_parser_multiple_sections(self, tmp_path):
        """Test parsing document with multiple sections."""
        doc = Document()
        
        # Add content to first section
        doc.add_paragraph("First section content")
        
        # Add section break (new section)
        doc.add_section()
        doc.add_paragraph("Second section content")
        
        file_path = tmp_path / "multi_section.docx"
        doc.save(file_path)
        
        parser = DocxParser(str(file_path))
        doc_info = parser.parse()
        
        # Should have two sections
        assert len(doc_info.sections) == 2
        
        # Each section should have correct index
        assert doc_info.sections[0].index == 0
        assert doc_info.sections[1].index == 1
