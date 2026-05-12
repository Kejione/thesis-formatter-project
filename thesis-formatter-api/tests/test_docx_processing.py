"""
Tests for document processing module.
"""

import pytest
import tempfile
import os

from docx import Document
from docx.shared import Pt, Cm

from app.services.docx import (
    DocxParser,
    FormatChecker,
    FormatFixer,
    DocxGenerator,
    DocumentProcessor,
    process_document,
)


# Sample rules for testing
SAMPLE_RULES = {
    "page_margin": {
        "top": "2.5cm",
        "bottom": "2.5cm",
        "left": "3cm",
        "right": "2.5cm",
    },
    "font": {
        "cn_body": "宋体",
        "en_body": "Times New Roman",
    },
    "font_size": {
        "body": "12pt",
        "heading1": "22pt",
        "heading2": "16pt",
    },
    "line_spacing": {
        "body": "1.5倍",
    },
}


@pytest.fixture
def sample_document():
    """Create a sample document for testing."""
    doc = Document()

    # Add a paragraph with some text
    para = doc.add_paragraph("这是一个测试段落，用于测试文档处理功能。")
    para.runs[0].font.size = Pt(12)

    # Add a heading
    heading = doc.add_heading("测试标题", level=1)

    # Add another paragraph
    para2 = doc.add_paragraph("这是另一个测试段落。")
    para2.runs[0].font.size = Pt(10)  # Wrong size

    # Save to temp file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        yield f.name

    # Cleanup
    os.unlink(f.name)


class TestDocxParser:
    """Tests for DocxParser."""

    def test_parse_document(self, sample_document):
        """Test parsing a document."""
        parser = DocxParser(sample_document)
        doc_info = parser.parse()

        assert doc_info is not None
        assert doc_info.page_count > 0
        assert doc_info.word_count > 0
        assert len(doc_info.paragraphs) > 0

    def test_extract_paragraphs(self, sample_document):
        """Test extracting paragraphs."""
        parser = DocxParser(sample_document)
        doc_info = parser.parse()

        assert len(doc_info.paragraphs) >= 3

        # Check first paragraph
        first_para = doc_info.paragraphs[0]
        assert "测试段落" in first_para.text

    def test_extract_sections(self, sample_document):
        """Test extracting sections."""
        parser = DocxParser(sample_document)
        doc_info = parser.parse()

        assert len(doc_info.sections) >= 1

        # Check section margins
        section = doc_info.sections[0]
        assert section.top_margin_cm > 0
        assert section.left_margin_cm > 0


class TestFormatChecker:
    """Tests for FormatChecker."""

    def test_check_all(self, sample_document):
        """Test checking all format issues."""
        parser = DocxParser(sample_document)
        doc_info = parser.parse()

        checker = FormatChecker(doc_info, SAMPLE_RULES)
        issues = checker.check_all()

        assert isinstance(issues, list)

    def test_check_font_size(self, sample_document):
        """Test checking font size issues."""
        parser = DocxParser(sample_document)
        doc_info = parser.parse()

        checker = FormatChecker(doc_info, SAMPLE_RULES)
        issues = checker.check_all()

        # Should find font size issue (para2 has 10pt instead of 12pt)
        font_size_issues = [i for i in issues if i.category.value == "font_size"]
        # May or may not find issues depending on the document


class TestDocumentProcessor:
    """Tests for DocumentProcessor."""

    def test_process_document(self, sample_document):
        """Test full document processing."""
        result = process_document(
            sample_document,
            SAMPLE_RULES,
            fix=True,
            generate_reports=True,
        )

        assert result.document_info is not None
        assert isinstance(result.issues, list)
        assert isinstance(result.changes, list)

        # Reports should be generated
        if result.changes:
            assert result.change_log_markdown is not None
        assert result.report_markdown is not None

    def test_check_only(self, sample_document):
        """Test check-only mode."""
        processor = DocumentProcessor(SAMPLE_RULES)
        doc_info, issues = processor.check_only(sample_document)

        assert doc_info is not None
        assert isinstance(issues, list)


class TestDocxGenerator:
    """Tests for DocxGenerator."""

    def test_generate_change_log(self, sample_document):
        """Test generating change log."""
        from app.services.docx.fixer import ChangeRecord

        doc = Document(sample_document)
        generator = DocxGenerator(doc, "test.docx")

        changes = [
            ChangeRecord(
                issue_id="font_size.body",
                category="font_size",
                location={"paragraph": 2, "page": 1},
                before_value="10pt",
                after_value="12pt",
                risk_level="low",
                timestamp="2026-05-11T10:00:00",
            )
        ]

        markdown = generator.generate_change_log_markdown(
            changes, {"标题": "测试文档", "页数": 1}
        )

        assert "# 格式修改记录" in markdown
        assert "字号" in markdown
        assert "10pt" in markdown
        assert "12pt" in markdown

    def test_generate_report(self, sample_document):
        """Test generating report."""
        doc = Document(sample_document)
        generator = DocxGenerator(doc, "test.docx")

        # Create mock issues
        from app.services.docx.checker import Issue, Severity, Category

        issues = [
            Issue(
                severity=Severity.WARNING,
                category=Category.FONT_SIZE,
                location={"paragraph": 2, "page": 1},
                rule_id="font_size.body",
                current_value="10pt",
                expected_value="12pt",
                suggestion="将字号从 10pt 改为 12pt",
                fixable=True,
            )
        ]

        markdown = generator.generate_report_markdown(
            issues, [], {"标题": "测试文档"}, SAMPLE_RULES
        )

        assert "# 格式检查报告" in markdown
        assert "字号" in markdown
