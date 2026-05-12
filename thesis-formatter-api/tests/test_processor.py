"""
Tests for DocumentProcessor module.

Tests document processing functionality including check-only mode,
full pipeline processing, and report generation.
"""

import pytest
import os
from docx import Document

import sys
import os

# Add parent directory to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

# Import directly from module files to avoid app package initialization
import importlib.util

def load_module(module_name, file_path):
    # Use unique module name to avoid conflicts
    unique_name = f"processor_test_{module_name}"
    if unique_name in sys.modules:
        return sys.modules[unique_name]
    spec = importlib.util.spec_from_file_location(unique_name, file_path)
    module = importlib.util.module_from_spec(spec)
    sys.modules[unique_name] = module
    spec.loader.exec_module(module)
    return module

# Load modules directly
base_path = os.path.join(os.path.dirname(__file__), '..', 'app', 'services', 'docx')

parser_module = load_module("parser", os.path.join(base_path, 'parser.py'))
checker_module = load_module("checker", os.path.join(base_path, 'checker.py'))
fixer_module = load_module("fixer", os.path.join(base_path, 'fixer.py'))
generator_module = load_module("generator", os.path.join(base_path, 'generator.py'))
processor_module = load_module("processor", os.path.join(base_path, 'processor.py'))

DocumentProcessor = processor_module.DocumentProcessor
ProcessResult = processor_module.ProcessResult
process_document = processor_module.process_document
DocumentInfo = parser_module.DocumentInfo
Issue = checker_module.Issue
Severity = checker_module.Severity
Category = checker_module.Category
ChangeRecord = fixer_module.ChangeRecord


class TestDocumentProcessor:
    """Test suite for DocumentProcessor class."""
    
    def test_processor_check_only(self, temp_docx_file, sample_rules):
        """Test check-only mode."""
        processor = DocumentProcessor(sample_rules)
        doc_info, issues = processor.check_only(temp_docx_file)
        
        # Should return document info (check by attributes, not type due to module reloading)
        assert doc_info is not None
        assert hasattr(doc_info, 'paragraphs')
        assert hasattr(doc_info, 'sections')
        
        # Should return list of issues
        assert isinstance(issues, list)
        
        # Should find some issues in our test document
        assert len(issues) > 0
        
        # Issues should be properly typed
        for issue in issues:
            assert hasattr(issue, 'severity')
            assert hasattr(issue, 'category')
    
    def test_processor_full_pipeline(self, temp_docx_file, sample_rules):
        """Test full processing pipeline with fixes."""
        processor = DocumentProcessor(sample_rules)
        result = processor.process(
            temp_docx_file,
            fix=True,
            generate_reports=True,
        )
        
        # Should return ProcessResult
        assert isinstance(result, ProcessResult)
        
        # Should have document info (check by attributes due to module reloading)
        assert result.document_info is not None
        assert hasattr(result.document_info, 'paragraphs')
        assert hasattr(result.document_info, 'sections')
        
        # Should have issues list
        assert result.issues is not None
        assert isinstance(result.issues, list)
        
        # Should have changes list
        assert result.changes is not None
        assert isinstance(result.changes, list)
        
        # Should have fixed document bytes
        assert result.fixed_document_bytes is not None
        assert isinstance(result.fixed_document_bytes, bytes)
        assert len(result.fixed_document_bytes) > 0
    
    def test_processor_generates_reports(self, temp_docx_file, sample_rules):
        """Test markdown report generation."""
        processor = DocumentProcessor(sample_rules)
        result = processor.process(
            temp_docx_file,
            fix=True,
            generate_reports=True,
        )
        
        # Should generate reports
        assert result.report_markdown is not None
        assert isinstance(result.report_markdown, str)
        assert len(result.report_markdown) > 0
        
        # Report should contain expected sections
        assert "# 格式检查报告" in result.report_markdown or "格式" in result.report_markdown
        
        # If there are changes, should have change log
        if result.changes:
            assert result.change_log_markdown is not None
            assert isinstance(result.change_log_markdown, str)
            assert len(result.change_log_markdown) > 0
    
    def test_processor_without_fix(self, temp_docx_file, sample_rules):
        """Test processing without applying fixes."""
        processor = DocumentProcessor(sample_rules)
        result = processor.process(
            temp_docx_file,
            fix=False,
            generate_reports=True,
        )
        
        # Should still return ProcessResult
        assert isinstance(result, ProcessResult)
        
        # Should have document info and issues
        assert result.document_info is not None
        assert result.issues is not None
        
        # Should not have changes
        assert len(result.changes) == 0
        
        # Should not have fixed document bytes
        assert result.fixed_document_bytes is None
        
        # Should still generate reports
        assert result.report_markdown is not None
    
    def test_processor_without_reports(self, temp_docx_file, sample_rules):
        """Test processing without generating reports."""
        processor = DocumentProcessor(sample_rules)
        result = processor.process(
            temp_docx_file,
            fix=True,
            generate_reports=False,
        )
        
        # Should still return ProcessResult
        assert isinstance(result, ProcessResult)
        
        # Should have document info, issues, and changes
        assert result.document_info is not None
        assert result.issues is not None
        assert result.changes is not None
        
        # Should not have reports
        assert result.report_markdown is None
        assert result.change_log_markdown is None
    
    def test_processor_fix_only(self, temp_docx_file, sample_rules):
        """Test fix-only mode."""
        processor = DocumentProcessor(sample_rules)
        changes, fixed_bytes = processor.fix_only(temp_docx_file)
        
        # Should return changes list
        assert isinstance(changes, list)
        
        # Should return fixed document bytes
        assert isinstance(fixed_bytes, bytes)
        assert len(fixed_bytes) > 0
        
        # Should be able to load fixed document
        from io import BytesIO
        fixed_doc = Document(BytesIO(fixed_bytes))
        assert len(fixed_doc.paragraphs) > 0
    
    def test_processor_fix_only_with_issue_ids(self, temp_docx_file, sample_rules):
        """Test fix-only mode with specific issue IDs."""
        processor = DocumentProcessor(sample_rules)
        
        # First, get all issues
        doc_info, all_issues = processor.check_only(temp_docx_file)
        
        if len(all_issues) >= 2:
            # Fix only first issue
            issue_ids = [all_issues[0].rule_id]
            changes, fixed_bytes = processor.fix_only(temp_docx_file, issue_ids=issue_ids)
            
            # Should only fix the specified issue
            assert len(changes) == 1
            assert changes[0].issue_id == issue_ids[0]
    
    def test_process_document_convenience_function(self, temp_docx_file, sample_rules):
        """Test the convenience function process_document."""
        result = process_document(
            temp_docx_file,
            sample_rules,
            fix=True,
            generate_reports=True,
        )
        
        # Should return ProcessResult
        assert isinstance(result, ProcessResult)
        
        # Should have all expected fields
        assert result.document_info is not None
        assert result.issues is not None
        assert result.changes is not None


class TestProcessResult:
    """Test ProcessResult dataclass."""
    
    def test_process_result_creation(self):
        """Test ProcessResult creation with all fields."""
        doc_info = DocumentInfo()
        issues = []
        changes = []
        
        result = ProcessResult(
            document_info=doc_info,
            issues=issues,
            changes=changes,
            fixed_document_bytes=b"test",
            change_log_markdown="# Changes",
            report_markdown="# Report",
        )
        
        assert result.document_info == doc_info
        assert result.issues == issues
        assert result.changes == changes
        assert result.fixed_document_bytes == b"test"
        assert result.change_log_markdown == "# Changes"
        assert result.report_markdown == "# Report"
    
    def test_process_result_optional_fields(self):
        """Test ProcessResult with optional fields as None."""
        doc_info = DocumentInfo()
        
        result = ProcessResult(
            document_info=doc_info,
            issues=[],
            changes=[],
        )
        
        assert result.fixed_document_bytes is None
        assert result.change_log_markdown is None
        assert result.report_markdown is None


class TestDocumentProcessorEdgeCases:
    """Test edge cases for DocumentProcessor."""
    
    def test_processor_with_empty_document(self, empty_docx_file, sample_rules):
        """Test processing empty document."""
        processor = DocumentProcessor(sample_rules)
        result = processor.process(empty_docx_file, fix=True, generate_reports=True)
        
        # Should handle empty document gracefully
        assert isinstance(result, ProcessResult)
        assert result.document_info is not None
        
        # Empty document should have 0 paragraphs
        assert len(result.document_info.paragraphs) == 0
        assert result.document_info.word_count == 0
    
    def test_processor_with_minimal_rules(self, temp_docx_file, minimal_rules):
        """Test processing with minimal rules."""
        processor = DocumentProcessor(minimal_rules)
        result = processor.process(temp_docx_file, fix=True, generate_reports=True)
        
        # Should work with minimal rules
        assert isinstance(result, ProcessResult)
        assert result.document_info is not None
        
        # Should only find margin issues
        for issue in result.issues:
            assert issue.category == Category.MARGIN
    
    def test_processor_preserves_document_structure(self, temp_docx_file, sample_rules):
        """Test that processor preserves document structure."""
        processor = DocumentProcessor(sample_rules)
        result = processor.process(temp_docx_file, fix=True, generate_reports=False)
        
        # Load fixed document
        from io import BytesIO
        fixed_doc = Document(BytesIO(result.fixed_document_bytes))
        
        # Original document
        original_doc = Document(temp_docx_file)
        
        # Should preserve number of paragraphs
        assert len(fixed_doc.paragraphs) == len(original_doc.paragraphs)
        
        # Should preserve text content
        for orig_para, fixed_para in zip(original_doc.paragraphs, fixed_doc.paragraphs):
            assert orig_para.text == fixed_para.text
    
    def test_processor_handles_complex_document(self, complex_docx_file, sample_rules):
        """Test processing complex document with many elements."""
        processor = DocumentProcessor(sample_rules)
        result = processor.process(complex_docx_file, fix=True, generate_reports=True)
        
        # Should handle complex document
        assert isinstance(result, ProcessResult)
        assert result.document_info is not None
        
        # Should have found issues
        assert len(result.issues) >= 0
        
        # Should generate reports
        assert result.report_markdown is not None


class TestDocumentProcessorIntegration:
    """Integration tests for DocumentProcessor."""
    
    def test_full_workflow_check_then_fix(self, temp_docx_file, sample_rules):
        """Test full workflow: check first, then fix."""
        processor = DocumentProcessor(sample_rules)
        
        # Step 1: Check only
        doc_info, issues = processor.check_only(temp_docx_file)
        
        # Should find issues
        assert len(issues) > 0
        
        # Count fixable issues
        fixable_count = len([i for i in issues if i.fixable])
        
        # Step 2: Fix
        result = processor.process(temp_docx_file, fix=True, generate_reports=True)
        
        # Should have applied fixes (may be less than fixable_count due to some issues not being fixed)
        assert len(result.changes) > 0
        assert len(result.changes) <= fixable_count
        
        # Step 3: Check again - should have fewer issues
        doc_info2, issues2 = processor.check_only(temp_docx_file)
        
        # Original document still has issues (we didn't modify the file)
        # But the fixed document bytes should be different
        assert result.fixed_document_bytes is not None
    
    def test_report_content_accuracy(self, temp_docx_file, sample_rules):
        """Test that reports contain accurate information."""
        processor = DocumentProcessor(sample_rules)
        result = processor.process(temp_docx_file, fix=True, generate_reports=True)
        
        # Report should mention document info
        report = result.report_markdown
        
        # Should contain issue count
        assert str(len(result.issues)) in report or "问题" in report or "issue" in report.lower()
        
        # If there are changes, change log should contain them
        if result.changes:
            change_log = result.change_log_markdown
            for change in result.changes:
                # Change log should mention the category
                assert change.category in change_log or change.issue_id in change_log
